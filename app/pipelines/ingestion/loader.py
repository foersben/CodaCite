"""Document loading and content extraction utilities.

This module provides a unified interface for extracting structured text content
from diverse file formats (PDF, DOCX, Markdown, Text). It implements a
high-fidelity extraction strategy using the Docling engine for layout-aware
PDF parsing and local Vision Language Models (VLM) for describing technical
schematics and imagery.
"""

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument

try:
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import (
        AcceleratorOptions,
        PdfPipelineOptions,
        RapidOcrOptions,
    )
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling_core.transforms.serializer.markdown import MarkdownDocSerializer
    from docling_core.types.doc.document import PictureItem, TableItem, TextItem
except ImportError:
    # Fallback for environments where docling is not yet installed
    DocumentConverter = None  # type: ignore
    InputFormat = None  # type: ignore
    PdfPipelineOptions = None  # type: ignore
    PdfFormatOption = None  # type: ignore
    MarkdownDocSerializer = None  # type: ignore
    RapidOcrOptions = None  # type: ignore
    PictureItem = None  # type: ignore
    TableItem = None  # type: ignore
    TextItem = None  # type: ignore

try:
    import torch
except ImportError:
    torch = None  # type: ignore

from app.pipelines.generation.vlm import LocalVLM

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Domain model representing a document's extracted state.

    This container encapsulates the normalized output of the loading stage,
    preserving the relationship between the raw content and its provenance.

    Attributes:
        text: The extracted and normalized text content (Markdown formatted for PDFs).
        source: The filesystem path or URI of the original document.
        format: The detected file format (e.g., 'pdf', 'docx', 'markdown').
    """

    text: str
    source: str
    format: str


class DocumentLoader:
    """Standardized orchestrator for multi-format document extraction.

    The DocumentLoader serves as the gateway to the ingestion pipeline,
    transforming heterogeneous binary formats into a uniform text representation.
    It leverages specialized extraction engines based on file extension and
    implements 'smart-routing' for hardware acceleration.

    Core Capabilities:
        - **Layout-Aware PDF Extraction**: Utilizes `Docling` to preserve tables,
          structural hierarchy, and semantic relationships.
        - **Visual Reasoning**: Integrates `LocalVLM` to generate textual
          descriptions of non-textual elements (images, technical drawings).
        - **Dynamic Acceleration**: Automatically selects the optimal compute
          device (CUDA, MPS, or CPU) based on available VRAM.

    Pipeline Position:
        Phase 1: Normalization. Converts `disk_path` -> `LoadedDocument` (Markdown).
    """

    _SUPPORTED_FORMATS: dict[str, str] = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
    }

    def __init__(self, vlm: LocalVLM | None = None) -> None:
        """Initialize the document loader.

        Args:
            vlm: Optional pre-initialized LocalVLM instance. If omitted,
                the loader will lazy-initialize a VLM upon encountering
                the first visual element.
        """
        self._vlm = vlm

    @property
    def vlm(self) -> LocalVLM:
        """Access the underlying Vision Language Model instance.

        Returns:
            A ready-to-use LocalVLM for image description.
        """
        if self._vlm is None:
            logger.info("[LOADER] Lazy-initializing LocalVLM for image processing")
            self._vlm = LocalVLM()
        return self._vlm

    def _get_optimal_device(self) -> str:
        """Determine the most efficient hardware accelerator for extraction.

        This method performs a runtime check of the system's hardware capabilities.
        It prioritizes NVIDIA GPUs (CUDA) if at least 1.5GB of VRAM is free,
        followed by Apple Silicon (MPS), and falls back to CPU as a safe baseline.

        Returns:
            The string identifier of the optimal device ('cuda', 'mps', or 'cpu').
        """
        if torch is None:
            return "cpu"

        # 1. Check NVIDIA CUDA
        if torch.cuda.is_available():
            try:
                free_mem, _ = torch.cuda.mem_get_info()
                free_gb = free_mem / (1024**3)
                logger.info("[LOADER] GPU VRAM check: %.2f GB free", free_gb)

                if free_gb >= 1.5:
                    return "cuda"
                else:
                    logger.warning(
                        "[LOADER] Insufficient VRAM for GPU extraction (Need 1.5GB, have %.2fGB). Falling back to CPU.",
                        free_gb,
                    )
                    return "cpu"
            except Exception as e:
                logger.error("[LOADER] Failed to check VRAM: %s. Falling back to CPU.", e)
                return "cpu"

        # 2. Check Apple Silicon (MPS)
        if hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
            logger.info("[LOADER] Apple Silicon (MPS) detected. Using MPS for extraction.")
            return "mps"

        # 3. Default Fallback
        return "cpu"

    def load(self, path: Path) -> list[LoadedDocument]:
        """Extract text from a document located at the given path.

        This is the primary entry point for document loading. It detects the
        file format via extension and routes the request to the appropriate
        internal extraction method.

        Args:
            path: Absolute filesystem path to the target document.

        Returns:
            A list containing the extracted LoadedDocument.

        Raises:
            ValueError: If the file extension does not match any supported format.
            RuntimeError: If extraction fails due to file corruption or engine error.
        """
        logger.info("[LOADER] Loading document from: %s", path)
        suffix = path.suffix.lower()
        fmt = self._SUPPORTED_FORMATS.get(suffix)

        if fmt is None:
            logger.error("[LOADER] Unsupported file format: %s", suffix)
            raise ValueError(
                f"Unsupported file format '{suffix}'. "
                f"Supported formats: {list(self._SUPPORTED_FORMATS)}"
            )

        try:
            if fmt == "pdf":
                text = self._load_pdf(path)
            elif fmt == "docx":
                text = self._load_docx(path)
            elif fmt == "text":
                text = self._load_text(path)
            else:
                text = self._load_markdown(path)
        except Exception as e:
            logger.error("[LOADER] Extraction failed for %s: %s", path, str(e))
            raise RuntimeError(f"Document extraction failed: {e}") from e

        logger.info("[LOADER] Successfully loaded %s document (%d characters)", fmt, len(text))
        return [LoadedDocument(text=text, source=str(path), format=fmt)]

    def _load_pdf(self, path: Path) -> str:
        """Extract text and structural elements from a PDF via Docling.

        This method orchestrates a complex extraction pipeline that includes:
        1. **Hybrid Parsing**: Layout analysis combined with OCR (RapidOCR).
        2. **Multi-Modal Synthesis**: Images are described by a LocalVLM and
           injected as Markdown blockquotes.
        3. **Structural Serialization**: Tables and lists are converted to
           standard Markdown for semantic preservation.

        Args:
            path: filesystem path to the PDF document.

        Returns:
            The complete document represented as a Markdown string.

        Raises:
            ImportError: If required extraction libraries (Docling) are missing.
            RuntimeError: If the document converter encounters a fatal error.
        """
        if DocumentConverter is None:
            logger.error("[LOADER] docling is not installed. PDF extraction failed.")
            raise ImportError("docling is not installed")

        logger.info("[LOADER] Converting PDF with Docling: %s", path)

        # Configure Docling with dynamic device selection and enable lightweight RapidOCR
        optimal_device = self._get_optimal_device()
        logger.info("[LOADER] Configuring Docling Accelerator for device: %s", optimal_device)

        pipeline_options = PdfPipelineOptions()
        pipeline_options.accelerator_options = AcceleratorOptions(
            num_threads=4, device=optimal_device
        )
        pipeline_options.do_ocr = True
        pipeline_options.ocr_options = RapidOcrOptions()
        pipeline_options.images_scale = 2.0
        pipeline_options.generate_page_images = False  # Memory optimization
        pipeline_options.generate_picture_images = False  # Memory optimization

        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
        )

        try:
            result = converter.convert(path)
            doc = result.document
        except Exception as e:
            logger.error("[LOADER] Docling conversion failed: %s", str(e))
            raise RuntimeError(f"Failed to parse document: {e}") from e

        if MarkdownDocSerializer is None:
            logger.error("[LOADER] docling_core is not installed. PDF extraction failed.")
            raise ImportError("docling_core is not installed")

        serializer = MarkdownDocSerializer(doc=doc)

        text_parts = []
        item_count = 0
        table_count = 0
        picture_count = 0

        logger.info("[LOADER] Starting item-by-item extraction from %s", path.name)

        for item, _level in doc.iterate_items():
            item_count += 1
            if isinstance(item, TableItem):
                table_count += 1
                logger.debug("[LOADER] Extracting table #%d", table_count)
                text_parts.append(serializer.serialize(item=item).text)
            elif isinstance(item, PictureItem):
                picture_count += 1
                page_no = item.prov[0].page_no if item.prov else "unknown"
                logger.info("[LOADER] Found image #%d in PDF (Page %s)", picture_count, page_no)

                try:
                    img_data = item.image
                    if img_data and img_data.pil_image:
                        logger.info("[LOADER] Triggering VLM for image #%d", picture_count)
                        description = self.vlm.describe_image(img_data.pil_image)
                        if description:
                            text_parts.append(
                                f"\n> [Technical Drawing Description: {description}]\n"
                            )
                            logger.info(
                                "[LOADER] VLM described image #%d: %s...",
                                picture_count,
                                description[:50],
                            )
                        else:
                            logger.warning(
                                "[LOADER] VLM returned empty description for image #%d",
                                picture_count,
                            )
                    else:
                        logger.debug(
                            "[LOADER] Picture item #%d has no image data to describe", picture_count
                        )
                except Exception as ve:
                    logger.error(
                        "[LOADER] VLM processing failed for image #%d: %s", picture_count, str(ve)
                    )
                    text_parts.append("\n[Image description failed]\n")
            else:
                # Normal text item
                content = serializer.serialize(item=item).text.strip()
                if content:
                    text_parts.append(content)

        logger.info(
            "[LOADER] PDF Extraction summary for %s: %d items, %d tables, %d images processed.",
            path.name,
            item_count,
            table_count,
            picture_count,
        )

        if not text_parts:
            logger.warning(
                "[LOADER] Extraction loop produced no text parts, falling back to full export."
            )
            return doc.export_to_markdown()

        return "\n\n".join(text_parts)

    @staticmethod
    def _load_docx(path: Path) -> str:
        """Extract text content from a Microsoft Word (DOCX) document.

        Args:
            path: Absolute filesystem path to the DOCX file.

        Returns:
            The concatenated text content of all paragraphs.
        """
        doc = DocxDocument(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        return "\n".join(paragraphs)

    @staticmethod
    def _load_markdown(path: Path) -> str:
        """Load and return the content of a Markdown file.

        Args:
            path: Absolute filesystem path to the Markdown file.

        Returns:
            The raw Markdown text.
        """
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _load_text(path: Path) -> str:
        """Load and return the content of a plain text file.

        Uses UTF-8 encoding with character replacement for robust loading of
        legacy or corrupted text files.

        Args:
            path: Absolute filesystem path to the text file.

        Returns:
            The raw text content.
        """
        return path.read_text(encoding="utf-8", errors="replace")
