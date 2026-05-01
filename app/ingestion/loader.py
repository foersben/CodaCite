"""Document loading utilities for various file formats.

This module provides a unified interface for loading text content from PDF,
DOCX, Markdown, and plain text files. It uses Docling for layout-aware PDF
extraction and a local VLM for describing technical drawings.
"""

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument

try:
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling_core.transforms.serializer.markdown import MarkdownDocSerializer
    from docling_core.types.doc.document import PictureItem, TableItem, TextItem
except ImportError:
    # Fallback for environments where docling is not yet installed
    DocumentConverter = None  # type: ignore
    InputFormat = None  # type: ignore
    PdfPipelineOptions = None  # type: ignore
    PdfFormatOption = None  # type: ignore
    MarkdownDocSerializer = None  # type: ignore
    PictureItem = None  # type: ignore
    TableItem = None  # type: ignore
    TextItem = None  # type: ignore

from app.infrastructure.vlm import LocalVLM

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Represents a document loaded from disk.

    Attributes:
        text: The raw text content of the document.
        source: The file path or source identifier.
        format: The file format (e.g., 'pdf', 'docx').
    """

    text: str
    source: str
    format: str


class DocumentLoader:
    """Standardized entry point for extracting text from diverse file formats.

    This class serves as the first stage of the ingestion pipeline. It maps file
    extensions to specific extraction engines (docling, python-docx, etc.) and
    normalizes the output into `LoadedDocument` objects.

    Supported Formats:
        - PDF: Extracted using `docling` (layout-aware, tables as MD).
        - DOCX: Extracted using `python-docx` (paragraph iteration).
        - Markdown/Text: Raw UTF-8 reading with error replacement.

    Pipeline Role:
        `UploadFile (Bytes)` -> `NamedTemporaryFile (Disk)` -> `DocumentLoader.load()` -> `list[LoadedDocument]`
    """

    _SUPPORTED_FORMATS: dict[str, str] = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
    }

    def __init__(self) -> None:
        """Initialize the loader with infrastructure adapters."""
        self.vlm = LocalVLM()

    def load(self, path: Path) -> list[LoadedDocument]:
        """Load a document from the specified path.

        Args:
            path: Path to the document file.

        Returns:
            A list containing a single LoadedDocument instance.

        Raises:
            ValueError: If the file format is not supported.
        """
        logger.info("[LOADER] Loading document from: %s", path)
        suffix = path.suffix.lower()
        fmt = self._SUPPORTED_FORMATS.get(suffix)

        if fmt is None:
            logger.error("[LOADER] Unsupported file format: %s", suffix)
            raise ValueError(
                f"Unsupported file format '{suffix}'. "
                f"Supported formats: {list(self._SUPPORTED_FORMATS)}"
            )

        if fmt == "pdf":
            text = self._load_pdf(path)
        elif fmt == "docx":
            text = self._load_docx(path)
        elif fmt == "text":
            text = self._load_text(path)
        else:
            text = self._load_markdown(path)

        logger.info("[LOADER] Successfully loaded %s document (%d characters)", fmt, len(text))
        return [LoadedDocument(text=text, source=str(path), format=fmt)]

    def _load_pdf(self, path: Path) -> str:
        """Extract text and layout from a PDF file using Docling.

        Args:
            path: Path to the PDF file.

        Returns:
            Extracted text content in Markdown format, with VLM descriptions
            for images and technical drawings.
        """
        if DocumentConverter is None:
            logger.error("[LOADER] docling is not installed. PDF extraction failed.")
            return "[Error: docling not installed]"

        logger.info("[LOADER] Converting PDF with Docling: %s", path)

        # Configure Docling to extract images for VLM processing
        pipeline_options = PdfPipelineOptions()
        pipeline_options.images_scale = 2.0  # High res for VLM
        pipeline_options.generate_page_images = True
        pipeline_options.generate_picture_images = True

        converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
        )

        try:
            result = converter.convert(path)
            doc = result.document
        except Exception as e:
            logger.error("[LOADER] Docling conversion failed: %s", str(e))
            return f"[Error during PDF conversion: {str(e)}]"

        if MarkdownDocSerializer is None:
            logger.error("[LOADER] docling_core is not installed. PDF extraction failed.")
            return "[Error: docling_core not installed]"

        serializer = MarkdownDocSerializer(doc=doc)

        text_parts = []
        item_count = 0
        table_count = 0
        picture_count = 0

        logger.info("[LOADER] Starting item-by-item extraction from %s", path.name)

        for item, _level in doc.iterate_items():
            item_count += 1
            if isinstance(item, TableItem):
                table_count += 1
                logger.debug("[LOADER] Extracting table #%d", table_count)
                text_parts.append(serializer.serialize(item=item).text)
            elif isinstance(item, PictureItem):
                picture_count += 1
                page_no = item.prov[0].page_no if item.prov else "unknown"
                logger.info("[LOADER] Found image #%d in PDF (Page %s)", picture_count, page_no)

                try:
                    img_data = item.image
                    if img_data and img_data.pil_image:
                        logger.info("[LOADER] Triggering VLM for image #%d", picture_count)
                        description = self.vlm.describe_image(img_data.pil_image)
                        if description:
                            text_parts.append(
                                f"\n> [Technical Drawing Description: {description}]\n"
                            )
                            logger.info(
                                "[LOADER] VLM described image #%d: %s...",
                                picture_count,
                                description[:50],
                            )
                        else:
                            logger.warning(
                                "[LOADER] VLM returned empty description for image #%d",
                                picture_count,
                            )
                    else:
                        logger.debug(
                            "[LOADER] Picture item #%d has no image data to describe", picture_count
                        )
                except Exception as ve:
                    logger.error(
                        "[LOADER] VLM processing failed for image #%d: %s", picture_count, str(ve)
                    )
                    text_parts.append("\n[Image description failed]\n")
            else:
                # Normal text item
                content = serializer.serialize(item=item).text.strip()
                if content:
                    text_parts.append(content)

        logger.info(
            "[LOADER] PDF Extraction summary for %s: %d items, %d tables, %d images processed.",
            path.name,
            item_count,
            table_count,
            picture_count,
        )

        if not text_parts:
            logger.warning(
                "[LOADER] Extraction loop produced no text parts, falling back to full export."
            )
            return doc.export_to_markdown()

        return "\n\n".join(text_parts)

    @staticmethod
    def _load_docx(path: Path) -> str:
        """Extract text from a DOCX file.

        Args:
            path: Path to the DOCX file.

        Returns:
            Extracted text content.
        """
        doc = DocxDocument(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        return "\n".join(paragraphs)

    @staticmethod
    def _load_markdown(path: Path) -> str:
        """Read content from a Markdown file.

        Args:
            path: Path to the Markdown file.

        Returns:
            File content as text.
        """
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _load_text(path: Path) -> str:
        """Read content from a plain text file.

        Args:
            path: Path to the text file.

        Returns:
            File content as text.
        """
        return path.read_text(encoding="utf-8", errors="replace")
